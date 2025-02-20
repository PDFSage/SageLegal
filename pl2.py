#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
import re
import pickle
import os
from collections import OrderedDict
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

class Lawsuit:
    def __init__(
        self,
        sections=None,
        exhibits=None,
        header=None,
        documents=None,
        case_information="",
        law_firm_information=""
    ):
        if sections is None:
            sections = OrderedDict()
        if exhibits is None:
            exhibits = OrderedDict()
        if header is None:
            header = OrderedDict()
        if documents is None:
            documents = OrderedDict()
        self.sections = OrderedDict(sections)
        self.exhibits = OrderedDict(exhibits)
        self.header = OrderedDict(header)
        self.documents = OrderedDict(documents)
        self.case_information = case_information
        self.law_firm_information = law_firm_information

    def __repr__(self):
        h = "\n".join([f"  {k}: {v}" for k, v in self.header.items()])
        s = "\n".join([f"  {k}: {v}" for k, v in self.sections.items()])
        e = []
        for ek, ev in self.exhibits.items():
            e.append(f"  {ek}: " + ", ".join([f"{k}={v}" for k, v in ev.items()]))
        e = "\n".join(e)
        d = []
        for dk, dv in self.documents.items():
            d.append(f"  {dk}: {dv}")
        d = "\n".join(d)
        return (
            "Lawsuit Object:\n\n"
            "CASE INFORMATION:\n"
            f"  {self.case_information}\n\n"
            "LAW FIRM INFORMATION:\n"
            f"  {self.law_firm_information}\n\n"
            "HEADER:\n"
            f"{h}\n\n"
            "SECTIONS:\n"
            f"{s}\n\n"
            "EXHIBITS:\n"
            f"{e}\n\n"
            "DOCUMENTS:\n"
            f"{d}\n"
        )

def is_line_all_caps(line_str):
    if not re.search(r'[A-Z]', line_str):
        return False
    return not re.search(r'[a-z]', line_str)

def wrap_text_to_lines(pdf_canvas, full_text, font_name, font_size, max_width):
    pdf_canvas.setFont(font_name, font_size)
    paragraphs = full_text.split('\n')
    all_lines = []
    for paragraph in paragraphs:
        words = paragraph.split()
        if not words:
            all_lines.append(("", False))
            continue
        current_line = ""
        for word in words:
            test_line = word if not current_line else (current_line + " " + word)
            if pdf_canvas.stringWidth(test_line, font_name, font_size) <= max_width:
                current_line = test_line
            else:
                all_lines.append((current_line, True))
                current_line = word
        if current_line:
            all_lines.append((current_line, False))
    return all_lines

def draw_firm_name_vertical_center(pdf_canvas, text, page_width, page_height):
    pdf_canvas.saveState()
    pdf_canvas.setFont("Helvetica-Bold", 10)
    text_width = pdf_canvas.stringWidth(text, "Helvetica-Bold", 10)
    x_pos = 0.2 * inch
    y_center = page_height / 2.0
    y_pos = y_center - (text_width / 2.0)
    pdf_canvas.translate(x_pos, y_pos)
    pdf_canvas.rotate(90)
    pdf_canvas.drawString(0, 0, text)
    pdf_canvas.restoreState()

def draw_checkbox_line(pdf_canvas, text, x, y):
    pdf_canvas.drawString(x, y, text)

def generate_cover_sheet_pdf(pdf_canvas, page_width, page_height):
    pdf_canvas.setFont("Helvetica", 12)
    left_margin = 1.0 * inch
    top_position = page_height - 1.0 * inch
    line_height = 18
    current_y = top_position
    pdf_canvas.setLineWidth(2)
    pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.0 * inch)
    pdf_canvas.drawString(left_margin, current_y, "1. COURT: KING COUNTY SUPERIOR COURT")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "2. CASE ASSIGNMENT AREA:   ☐ Kent     ☒ Seattle")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "3. CASE TITLE: Bo Shang v. Amazon.com, Inc.")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "4. CASE NUMBER (Clerk to Assign): ______________________")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "5. CASE CATEGORY (Check one):   ☒ Civil")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "6. CASE TYPE:   ☒ TTO – Tort/Other")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "7. DOCUMENT/S BEING FILED:")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☐ Initial Pleadings and Petitions")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☐ Additional/Amended Pleadings")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☒ Complaint for Tort – Other (CMPTTO)")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☒ Summons")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "8. RELIEF REQUESTED:")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☒ Damages")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☒ Injunctive Relief")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☐ Other: ____________________")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "9. JURY DEMAND:")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☒ Yes")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   ☐ No")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "10. ATTORNEY OR PARTY SIGNING COVER SHEET:")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   Name:      Bo Shang (Plaintiff Pro Se)")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   Address:   10 McCafferty Way")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "              Burlington, MA 01803-3127")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   Phone:     781-999-4101")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   Email:     enigmatictyphoon@gmail.com")
    current_y -= line_height
    pdf_canvas.drawString(left_margin, current_y, "   WSBA No.:  Pro Se")

def generate_cover_sheet_docx(doc: Document):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    def add_line(text, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.bold = bold

    add_line("1. COURT: KING COUNTY SUPERIOR COURT")
    add_line("2. CASE ASSIGNMENT AREA:   ☐ Kent     ☒ Seattle")
    add_line("3. CASE TITLE: Bo Shang v. Amazon.com, Inc.")
    add_line("4. CASE NUMBER (Clerk to Assign): ______________________")
    add_line("5. CASE CATEGORY (Check one):   ☒ Civil")
    add_line("6. CASE TYPE:   ☒ TTO – Tort/Other")
    add_line("7. DOCUMENT/S BEING FILED:")
    add_line("   ☐ Initial Pleadings and Petitions")
    add_line("   ☐ Additional/Amended Pleadings")
    add_line("   ☒ Complaint for Tort – Other (CMPTTO)")
    add_line("   ☒ Summons")
    add_line("8. RELIEF REQUESTED:")
    add_line("   ☒ Damages")
    add_line("   ☒ Injunctive Relief")
    add_line("   ☐ Other: ____________________")
    add_line("9. JURY DEMAND:")
    add_line("   ☒ Yes")
    add_line("   ☐ No")
    add_line("10. ATTORNEY OR PARTY SIGNING COVER SHEET:")
    add_line("   Name:      Bo Shang (Plaintiff Pro Se)")
    add_line("   Address:   10 McCafferty Way")
    add_line("              Burlington, MA 01803-3127")
    add_line("   Phone:     781-999-4101")
    add_line("   Email:     enigmatictyphoon@gmail.com")
    add_line("   WSBA No.:  Pro Se")

def draw_exhibit_page(
    pdf_canvas,
    page_width,
    page_height,
    firm_name,
    case_name,
    exhibit_caption,
    exhibit_image,
    page_number,
    total_pages
):
    pdf_canvas.setLineWidth(2)
    pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.3 * inch)
    draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
    pdf_canvas.setFont("Helvetica-Bold", 12)
    pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
    pdf_canvas.setLineWidth(1)
    pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
    pdf_canvas.setFont("Helvetica", 10)
    top_margin = page_height - 1.2 * inch
    left_margin = 1.2 * inch
    line_spacing = 0.25 * inch
    max_caption_width = page_width - (2 * left_margin)
    from reportlab.pdfgen import canvas as dummy
    temp_c = dummy.Canvas("dummy", pagesize=letter)
    wrapped_caption_lines = wrap_text_to_lines(temp_c, exhibit_caption, "Helvetica", 10, max_caption_width)
    current_y = top_margin
    for (cap_line, _) in wrapped_caption_lines:
        pdf_canvas.drawString(left_margin, current_y, cap_line)
        current_y -= line_spacing
    top_of_image_area = current_y - line_spacing
    bottom_of_image_area = 0.5 * inch
    if top_of_image_area < bottom_of_image_area:
        top_of_image_area = bottom_of_image_area
    available_height = top_of_image_area - bottom_of_image_area
    available_width = page_width - 1.0 * inch
    try:
        img_reader = ImageReader(exhibit_image)
        img_width, img_height = img_reader.getSize()
    except Exception as e:
        pdf_canvas.setFont("Helvetica-Oblique", 10)
        pdf_canvas.drawCentredString(
            page_width / 2.0,
            page_height / 2.0,
            f"Unable to load image: {exhibit_image} Error: {e}"
        )
    else:
        scale = min(available_width / img_width, available_height / img_height, 1.0)
        new_width = img_width * scale
        new_height = img_height * scale
        x_img = 0.5 * inch + (available_width - new_width) / 2.0
        y_img_bottom = bottom_of_image_area
        pdf_canvas.drawImage(
            img_reader,
            x_img,
            y_img_bottom,
            width=new_width,
            height=new_height,
            preserveAspectRatio=True,
            anchor='c'
        )
    pdf_canvas.setFont("Helvetica-Oblique", 9)
    footer_text = f"Page {page_number} of {total_pages}"
    pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)

def is_full_equals_line(line_str):
    stripped = line_str.strip()
    if len(stripped) < 5:
        return False
    return bool(re.match(r'^[=]+$', stripped))

def detect_legal_title_blocks(lines):
    i = 0
    n = len(lines)
    while i < n:
        if is_full_equals_line(lines[i]):
            j = i + 1
            inner_lines = []
            found_bottom = False
            while j < n:
                if is_full_equals_line(lines[j]):
                    found_bottom = True
                    j += 1
                    break
                else:
                    inner_lines.append(lines[j])
                j += 1
            if found_bottom:
                yield ("legal_page_title_block", inner_lines)
                i = j
            else:
                yield ("normal_line", lines[i])
                i += 1
        else:
            yield ("normal_line", lines[i])
            i += 1

def draw_legal_page_title_block(
    pdf_canvas,
    page_width,
    page_height,
    block_lines,
    firm_name,
    case_name,
    page_number,
    total_pages
):
    pdf_canvas.setLineWidth(2)
    pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.3 * inch)
    draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
    pdf_canvas.setFont("Helvetica-Bold", 12)
    pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
    pdf_canvas.setLineWidth(1)
    pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
    pdf_canvas.setFont("Helvetica-Bold", 14)
    line_spacing = 0.3 * inch
    y_text = page_height - 1.5 * inch
    for line_str in block_lines:
        pdf_canvas.drawCentredString(page_width / 2.0, y_text, line_str)
        y_text -= line_spacing
    pdf_canvas.setFont("Helvetica-Oblique", 9)
    footer_text = f"Page {page_number} of {total_pages}"
    pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)

############################################################################
# Completely changed heuristics to detect references to case law/statutes
############################################################################
def extract_references(text):
    pattern = re.compile(
        r'(Barnes v\. Yahoo!.*?\(9th Cir\..*?\))|'
        r'(Cal\.?\s*Civ\.?\s*Code\s*§+\s*\d+(?:[–\-]\d+)*)|'
        r'(Lazar v\. Superior Court)|'
        r'(Fair Housing Council.*Roommates\.com.*?\(9th Cir\..*?\))|'
        r'(FTC v\. Accusearch.*?\(10th Cir\..*?\))',
        re.IGNORECASE
    )
    matches = pattern.findall(text)
    found = []
    for match in matches:
        for group in match:
            if group.strip():
                found.append(group.strip())
    return found

def draw_page_of_segments(
    pdf_canvas,
    page_width,
    page_height,
    segments,
    start_index,
    max_lines_per_page,
    firm_name,
    case_name,
    page_number,
    total_pages,
    line_offset_x,
    line_offset_y,
    line_spacing,
    heading_positions,
    reference_positions
):
    pdf_canvas.setLineWidth(2)
    pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.3 * inch)
    draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
    pdf_canvas.setFont("Helvetica-Bold", 12)
    pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
    pdf_canvas.setLineWidth(1)
    pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
    end_index = start_index
    current_line_count = 0
    y_text = line_offset_y
    while end_index < len(segments) and current_line_count < max_lines_per_page:
        seg = segments[end_index]
        if seg.get("page_always_new"):
            if current_line_count > 0:
                break
            else:
                block_lines = seg["lines"]
                draw_legal_page_title_block(
                    pdf_canvas,
                    page_width,
                    page_height,
                    block_lines,
                    firm_name,
                    case_name,
                    page_number,
                    total_pages
                )
                end_index += 1
                return end_index
        line_number = end_index + 1
        pdf_canvas.setFont("Helvetica", 10)
        pdf_canvas.drawString(line_offset_x - 0.6 * inch, y_text, str(line_number))
        pdf_canvas.drawString(page_width - 0.4 * inch, y_text, str(line_number))
        if seg["is_heading"] or seg["is_subheading"]:
            heading_positions.append((seg["text"], page_number, line_number, seg["is_subheading"]))
        text_line = seg["text"]
        references_found = extract_references(text_line)
        for ref in references_found:
            reference_positions.append((ref, page_number, line_number))
        pdf_canvas.setFont(seg["font_name"], seg["font_size"])
        if seg["alignment"] == "center":
            left_boundary = line_offset_x
            right_boundary = page_width - 0.5 * inch
            mid_x = (left_boundary + right_boundary) / 2.0
            pdf_canvas.drawCentredString(mid_x, y_text, text_line)
        else:
            pdf_canvas.drawString(line_offset_x, y_text, text_line)
        y_text -= line_spacing
        current_line_count += 1
        end_index += 1
    pdf_canvas.setFont("Helvetica-Oblique", 9)
    footer_text = f"Page {page_number} of {total_pages}"
    pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)
    return end_index

def generate_index_pdf(index_filename, firm_name, case_name, heading_positions, reference_positions):
    pdf_canvas = canvas.Canvas(index_filename, pagesize=letter)
    pdf_canvas.setTitle("Table of Contents")
    page_width, page_height = letter
    top_margin = 1.0 * inch
    bottom_margin = 1.0 * inch
    left_margin = 1.0 * inch
    right_margin = 0.5 * inch
    line_spacing = 0.25 * inch
    from reportlab.pdfgen import canvas as dummy
    temp_c = dummy.Canvas("dummy.pdf", pagesize=letter)

    def wrap_text(linestr, font_name, font_size, maxwidth):
        temp_c.setFont(font_name, font_size)
        wrapped = []
        paragraphs = linestr.split('\n')
        for paragraph in paragraphs:
            subwrap = wrap_text_to_lines(temp_c, paragraph, font_name, font_size, maxwidth)
            for (l, _) in subwrap:
                wrapped.append(l)
        return wrapped

    max_entry_width = page_width - left_margin - 1.5 * inch
    flattened_headings = []
    for (heading_text, pg_num, ln_num, is_sub) in heading_positions:
        if is_sub:
            f_name = "Helvetica"
            f_size = 9
            bld = False
        else:
            f_name = "Helvetica-Bold"
            f_size = 10
            bld = True
        lines = wrap_text(heading_text, f_name, f_size, max_entry_width)
        for i, txt_line in enumerate(lines):
            flattened_headings.append((txt_line, pg_num, ln_num, f_name, f_size, bld, (i==0)))
    flattened_refs = []
    for (ref_text, pg, ln) in reference_positions:
        f_name = "Helvetica"
        f_size = 9
        lines = wrap_text(ref_text, f_name, f_size, max_entry_width)
        for i, txt_line in enumerate(lines):
            flattened_refs.append((txt_line, pg, ln, f_name, f_size, False, (i==0)))
    all_index_entries = [('Heading', x) for x in flattened_headings] + [('Reference', x) for x in flattened_refs]
    usable_height = page_height - (top_margin + bottom_margin) - 1.0 * inch
    max_lines_per_page = int(usable_height // line_spacing)
    total_lines = len(all_index_entries)
    total_index_pages = max(1, (total_lines + max_lines_per_page - 1) // max_lines_per_page)
    i = 0
    current_page_index = 1
    while i < total_lines:
        pdf_canvas.setLineWidth(2)
        pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.3 * inch)
        draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
        pdf_canvas.setFont("Helvetica-Bold", 12)
        pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
        pdf_canvas.setLineWidth(1)
        pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
        pdf_canvas.setFont("Helvetica-Bold", 14)
        pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.75 * inch, "TABLE OF CONTENTS")
        x_text = left_margin
        y_text = page_height - top_margin - 0.75 * inch
        lines_on_this_page = 0
        while i < total_lines and lines_on_this_page < max_lines_per_page:
            etype, data = all_index_entries[i]
            (txt_line, pg_num, ln_num, fn, fs, bold, show_pg) = data
            pdf_canvas.setFont(fn, fs)
            pdf_canvas.drawString(x_text, y_text, txt_line)
            if show_pg:
                lbl = f"{pg_num}:{ln_num}"
                pdf_canvas.drawRightString(page_width - right_margin - 0.2 * inch, y_text, lbl)
            y_text -= line_spacing
            i += 1
            lines_on_this_page += 1
        pdf_canvas.setFont("Helvetica-Oblique", 9)
        footer_text = f"Index Page {current_page_index} of {total_index_pages}"
        pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)
        if i < total_lines:
            pdf_canvas.showPage()
            current_page_index += 1
        else:
            break
    pdf_canvas.save()

def generate_complaint_docx(docx_filename, firm_name, case_name, header_od, sections_od, heading_styles):
    doc = Document()
    generate_cover_sheet_docx(doc)
    doc.add_page_break()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    top_par = doc.add_paragraph()
    top_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top_par.add_run(f"{firm_name} | {case_name}\n")
    run.bold = True
    run.font.size = Pt(14)

    def add_center_big(line):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line.strip())
        r.bold = True
        r.font.size = Pt(14)

    header_content = header_od.get("content", "")
    header_lines = header_content.splitlines()
    buffer_of_lines = []
    for kind, block_lines in detect_legal_title_blocks(header_lines):
        if kind == "legal_page_title_block":
            if buffer_of_lines:
                for line in buffer_of_lines:
                    p = doc.add_paragraph()
                    ls = line.strip()
                    if is_line_all_caps(ls):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.add_run(ls)
                buffer_of_lines = []
            for line in block_lines:
                add_center_big(line)
        else:
            buffer_of_lines.append(block_lines)
    if buffer_of_lines:
        for line in buffer_of_lines:
            p = doc.add_paragraph()
            ls = line.strip()
            if is_line_all_caps(ls):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(ls)
    for section_key, section_body in sections_od.items():
        style_type = heading_styles.get(section_key, "section")
        doc.add_paragraph()
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if style_type == "section":
            run = heading_para.add_run(section_key)
            run.bold = True
            run.font.size = Pt(12)
        else:
            run = heading_para.add_run(section_key)
            run.bold = False
            run.font.size = Pt(11)
        body_lines = section_body.splitlines()
        normal_buffer = []
        def flush_buffer():
            for bline in normal_buffer:
                bls = bline.strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                rr = p.add_run(bls)
                if style_type == "section":
                    rr.font.size = Pt(12)
                else:
                    rr.font.size = Pt(11)
            normal_buffer.clear()
        for kind, block_lines in detect_legal_title_blocks(body_lines):
            if kind == "legal_page_title_block":
                flush_buffer()
                for xline in block_lines:
                    add_center_big(xline)
            else:
                normal_buffer.append(block_lines)
        flush_buffer()
    doc.save(docx_filename)

def generate_toc_docx(docx_filename, firm_name, case_name, heading_positions, reference_positions):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    top_par = doc.add_paragraph()
    top_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = top_par.add_run(f"{firm_name} | {case_name}\nTABLE OF CONTENTS\n")
    r.bold = True
    r.font.size = Pt(14)
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for (htext, pnum, lnum, is_sub) in heading_positions:
        if is_sub:
            fs = 11
            b = False
        else:
            fs = 12
            b = True
        row_cells = table.add_row().cells
        lp = row_cells[0].paragraphs[0]
        rp = row_cells[1].paragraphs[0]
        rleft = lp.add_run(htext)
        rleft.font.size = Pt(fs)
        rleft.bold = b
        rright = rp.add_run(f"{pnum}:{lnum}")
        rright.font.size = Pt(fs)
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for (ref_text, pnum, lnum) in reference_positions:
        row_cells = table.add_row().cells
        lp = row_cells[0].paragraphs[0]
        rp = row_cells[1].paragraphs[0]
        run_l = lp.add_run(ref_text)
        run_l.font.size = Pt(11)
        run_l.bold = False
        run_r = rp.add_run(f"{pnum}:{lnum}")
        run_r.font.size = Pt(11)
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(docx_filename)

def parse_documents_from_text(raw_text):
    lines = raw_text.splitlines()
    docs = []
    i = 0
    n = len(lines)
    while i < n:
        if is_full_equals_line(lines[i]):
            j = i + 1
            doc_lines = []
            while j < n and not is_full_equals_line(lines[j]):
                doc_lines.append(lines[j])
                j += 1
            if j < n:
                docs.append("\n".join(doc_lines))
                i = j + 1
            else:
                break
        else:
            i += 1
    return docs

def parse_header_and_sections(raw_text):
    header_od = OrderedDict()
    sections_od = OrderedDict()
    heading_pattern = re.compile(r'^((?:[IVXLCDM]+\.|[0-9]+\.)+)\s+(.*)$', re.IGNORECASE)
    lines = raw_text.splitlines()
    idx = 0
    header_lines = []
    while idx < len(lines):
        line = lines[idx].rstrip('\n').rstrip('\r')
        m = heading_pattern.match(line)
        if m:
            heading_number = m.group(1).strip()
            heading_title = m.group(2).strip()
            if is_line_all_caps(heading_title):
                break
        header_lines.append(line)
        idx += 1
    header_od["content"] = "\n".join(header_lines)
    current_heading_key = None
    current_body_lines = []
    while idx < len(lines):
        line = lines[idx].rstrip('\n').rstrip('\r')
        m = heading_pattern.match(line)
        if m:
            heading_number = m.group(1).strip()
            heading_title = m.group(2).strip()
            if is_line_all_caps(heading_title):
                if current_heading_key is not None:
                    sections_od[current_heading_key] = "\n".join(current_body_lines)
                current_body_lines = []
                if heading_number.endswith('.'):
                    heading_number = heading_number[:-1]
                current_heading_key = f"{heading_number} {heading_title}"
            else:
                current_body_lines.append(line)
        else:
            current_body_lines.append(line)
        idx += 1
    if current_heading_key is not None:
        sections_od[current_heading_key] = "\n".join(current_body_lines)
    return header_od, sections_od

def classify_headings(sections_od):
    heading_styles = {}
    for full_key in sections_od.keys():
        parts = full_key.split(None, 1)
        if len(parts) == 2:
            heading_number, _heading_text = parts[0], parts[1]
        else:
            heading_number = parts[0]
        dot_count = heading_number.count('.')
        if dot_count > 1:
            heading_styles[full_key] = "subsection"
        else:
            heading_styles[full_key] = "section"
    return heading_styles

def prepare_main_pdf_segments(header_text, sections_od, heading_styles, pdf_canvas, max_text_width):
    segments = []
    header_lines = header_text.splitlines()
    normal_buffer = []
    def flush_normal_buffer():
        for line in normal_buffer:
            ls = line.strip()
            if not ls:
                segments.append({
                    "text": "",
                    "font_name": "Helvetica",
                    "font_size": 10,
                    "alignment": "left",
                    "is_heading": False,
                    "is_subheading": False
                })
            elif is_line_all_caps(ls):
                wrapped = wrap_text_to_lines(pdf_canvas, ls, "Helvetica", 10, max_text_width)
                for (wl, _) in wrapped:
                    segments.append({
                        "text": wl,
                        "font_name": "Helvetica",
                        "font_size": 10,
                        "alignment": "center",
                        "is_heading": False,
                        "is_subheading": False
                    })
            else:
                wrapped = wrap_text_to_lines(pdf_canvas, ls, "Helvetica", 10, max_text_width)
                for (wl, _) in wrapped:
                    segments.append({
                        "text": wl,
                        "font_name": "Helvetica",
                        "font_size": 10,
                        "alignment": "left",
                        "is_heading": False,
                        "is_subheading": False
                    })
        normal_buffer.clear()
    for kind, block_lines in detect_legal_title_blocks(header_lines):
        if kind == "legal_page_title_block":
            flush_normal_buffer()
            lines_cleaned = [ln.strip() for ln in block_lines]
            segments.append({
                "legal_page_title": True,
                "page_always_new": True,
                "lines": lines_cleaned
            })
        else:
            normal_buffer.append(block_lines)
    flush_normal_buffer()

    for section_key, section_body in sections_od.items():
        style = heading_styles.get(section_key, "section")
        if style == "section":
            hfn = "Helvetica-Bold"
            hfs = 10
            bfn = "Helvetica"
            bfs = 10
            ish = True
            issub = False
        else:
            hfn = "Helvetica"
            hfs = 9
            bfn = "Helvetica"
            bfs = 9
            ish = False
            issub = True
        segments.append({
            "text": "",
            "font_name": bfn,
            "font_size": bfs,
            "alignment": "left",
            "is_heading": False,
            "is_subheading": False
        })
        heading_wrapped = wrap_text_to_lines(pdf_canvas, section_key, hfn, hfs, max_text_width)
        for (wl, _) in heading_wrapped:
            segments.append({
                "text": wl,
                "font_name": hfn,
                "font_size": hfs,
                "alignment": "center",
                "is_heading": ish,
                "is_subheading": issub
            })
        lines_of_body = section_body.splitlines()
        normal_buffer_sec = []
        def flush_section_buf():
            for line in normal_buffer_sec:
                ls = line.strip()
                if not ls:
                    segments.append({
                        "text": "",
                        "font_name": bfn,
                        "font_size": bfs,
                        "alignment": "left",
                        "is_heading": False,
                        "is_subheading": False
                    })
                else:
                    w = wrap_text_to_lines(pdf_canvas, ls, bfn, bfs, max_text_width)
                    for (wl, _) in w:
                        segments.append({
                            "text": wl,
                            "font_name": bfn,
                            "font_size": bfs,
                            "alignment": "left",
                            "is_heading": False,
                            "is_subheading": False
                        })
            normal_buffer_sec.clear()
        for kind, block_lines in detect_legal_title_blocks(lines_of_body):
            if kind == "legal_page_title_block":
                flush_section_buf()
                lines_cleaned = [ln.strip() for ln in block_lines]
                segments.append({
                    "legal_page_title": True,
                    "page_always_new": True,
                    "lines": lines_cleaned
                })
            else:
                normal_buffer_sec.append(block_lines)
        flush_section_buf()
    return segments

def generate_legal_document(
    firm_name,
    case_name,
    output_filename,
    header_od,
    sections_od,
    exhibits,
    heading_positions
):
    page_width, page_height = letter
    pdf_canvas = canvas.Canvas(output_filename, pagesize=letter)
    pdf_canvas.setTitle("Legal Document with Cover Sheet")
    pdf_canvas.setAuthor(firm_name)
    pdf_canvas.setSubject(case_name)
    pdf_canvas.setCreator("Legal PDF Generator")
    generate_cover_sheet_pdf(pdf_canvas, page_width, page_height)
    pdf_canvas.showPage()
    heading_styles = classify_headings(sections_od)
    top_margin = 1.0 * inch
    bottom_margin = 1.0 * inch
    left_margin = 1.2 * inch
    right_margin = 0.5 * inch
    line_spacing = 0.25 * inch
    usable_height = page_height - (top_margin + bottom_margin)
    max_lines_per_page = int(usable_height // line_spacing)
    line_offset_x = left_margin
    line_offset_y = page_height - top_margin
    max_text_width = page_width - right_margin - line_offset_x - 0.2 * inch
    segments = prepare_main_pdf_segments(
        header_text=header_od.get("content", ""),
        sections_od=sections_od,
        heading_styles=heading_styles,
        pdf_canvas=pdf_canvas,
        max_text_width=max_text_width
    )
    current_index = 0
    text_pages = 0
    total_segments = len(segments)
    while current_index < total_segments:
        seg = segments[current_index]
        if seg.get("page_always_new"):
            text_pages += 1
            current_index += 1
        else:
            lines_used = 0
            local_i = current_index
            while local_i < total_segments and lines_used < max_lines_per_page:
                s = segments[local_i]
                if s.get("page_always_new"):
                    break
                lines_used += 1
                local_i += 1
            text_pages += 1
            current_index = local_i
    exhibit_pages = len(exhibits)
    total_pages = 1 + text_pages + exhibit_pages
    page_number = 2
    current_index = 0
    reference_positions = []
    while current_index < total_segments:
        next_index = draw_page_of_segments(
            pdf_canvas=pdf_canvas,
            page_width=page_width,
            page_height=page_height,
            segments=segments,
            start_index=current_index,
            max_lines_per_page=max_lines_per_page,
            firm_name=firm_name,
            case_name=case_name,
            page_number=page_number,
            total_pages=total_pages,
            line_offset_x=line_offset_x,
            line_offset_y=line_offset_y,
            line_spacing=line_spacing,
            heading_positions=heading_positions,
            reference_positions=reference_positions
        )
        pdf_canvas.showPage()
        page_number += 1
        current_index = next_index
    for (caption, image_path) in exhibits:
        draw_exhibit_page(
            pdf_canvas=pdf_canvas,
            page_width=page_width,
            page_height=page_height,
            firm_name=firm_name,
            case_name=case_name,
            exhibit_caption=caption,
            exhibit_image=image_path,
            page_number=page_number,
            total_pages=total_pages
        )
        pdf_canvas.showPage()
        page_number += 1
    pdf_canvas.save()
    generate_complaint_docx(
        docx_filename=os.path.splitext(output_filename)[0] + ".docx",
        firm_name=firm_name,
        case_name=case_name,
        header_od=header_od,
        sections_od=sections_od,
        heading_styles=heading_styles
    )
    return reference_positions

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--firm_name", required=True)
    parser.add_argument("--case", required=True)
    parser.add_argument("--output", default="lawsuit.pdf")
    parser.add_argument("--file", required=True)
    parser.add_argument("--exhibits", nargs='+', default=[])
    parser.add_argument("--index", default="index.pdf")
    parser.add_argument("--pickle", nargs='?', const=None)
    args = parser.parse_args()
    with open(args.file, 'r', encoding='utf-8') as f:
        raw_text = f.read()
    header_od, sections_od = parse_header_and_sections(raw_text)
    if len(args.exhibits) % 2 != 0:
        raise ValueError("Exhibits must be in pairs: caption_file image_file")
    exhibits_od = OrderedDict()
    ex_index = 1
    for i in range(0, len(args.exhibits), 2):
        cap_file = args.exhibits[i]
        img_file = args.exhibits[i + 1]
        with open(cap_file, 'r', encoding='utf-8') as cfp:
            cap_text = cfp.read()
        exhibits_od[str(ex_index)] = OrderedDict([
            ('caption', cap_text),
            ('image_path', img_file)
        ])
        ex_index += 1
    header_od["DocumentTitle"] = "Complaint for Tort – Other"
    header_od["DateFiled"] = "2025-02-14"
    header_od["Court"] = "King County Superior Court"
    found_documents = parse_documents_from_text(raw_text)
    documents_od = OrderedDict()
    for idx, doc_text in enumerate(found_documents, start=1):
        documents_od[str(idx)] = doc_text
    lawsuit_obj = Lawsuit(
        sections=sections_od,
        exhibits=exhibits_od,
        header=header_od,
        documents=documents_od,
        case_information=args.case,
        law_firm_information=args.firm_name
    )
    exhibits_for_pdf = []
    for _, val in lawsuit_obj.exhibits.items():
        exhibits_for_pdf.append((val["caption"], val["image_path"]))
    heading_positions = []
    reference_positions = generate_legal_document(
        firm_name=args.firm_name,
        case_name=args.case,
        output_filename=args.output,
        header_od=header_od,
        sections_od=sections_od,
        exhibits=exhibits_for_pdf,
        heading_positions=heading_positions
    )
    generate_index_pdf(
        index_filename=args.index,
        firm_name=args.firm_name,
        case_name=args.case,
        heading_positions=heading_positions,
        reference_positions=reference_positions
    )
    index_docx = os.path.splitext(args.index)[0] + ".docx"
    generate_toc_docx(
        docx_filename=index_docx,
        firm_name=args.firm_name,
        case_name=args.case,
        heading_positions=heading_positions,
        reference_positions=reference_positions
    )
    if args.pickle is not None:
        pklfile = args.pickle if args.pickle else "lawsuit.pickle"
        with open(pklfile, "wb") as pf:
            pickle.dump(lawsuit_obj, pf)
    print(f"PDF generated: {args.output}")
    print(f"DOCX Complaint: {os.path.splitext(args.output)[0] + '.docx'}")
    print(f"Index PDF: {args.index}")
    print(f"Index DOCX: {index_docx}")
    print(lawsuit_obj)

if __name__ == "__main__":
    main()